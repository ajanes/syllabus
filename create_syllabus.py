#!/usr/bin/env python

import os
import re
import yaml
import argparse
import subprocess
from docx import Document
from pybtex.database.input import bibtex
from pybtex.plugin import find_plugin
from collections.abc import Mapping


def load_yaml(files):
    merged = {}
    for file in files:
        with open(file) as f:
            data = yaml.safe_load(f)
            if data:
                merge_dicts(merged, data)
                includes = data.get("course", {}).get("include", [])
                for inc in includes:
                    with open(os.path.abspath(inc)) as inc_file:
                        inc_data = yaml.safe_load(inc_file)
                        if inc_data:
                            merge_dicts(merged, inc_data)
    return merged


def merge_dicts(target, source):
    for key, value in source.items():
        if isinstance(value, Mapping) and key in target:
            merge_dicts(target[key], value)
        else:
            target[key] = value


def get_bib_entry(path, key):
    parser = bibtex.Parser()
    bib = parser.parse_file(path)
    if key not in bib.entries:
        return f"BibTeX key '{key}' not found."
    entry = bib.entries[key]
    style = find_plugin("pybtex.style.formatting", "plain")()
    backend = find_plugin("pybtex.backends", "text")()
    return next(style.format_entries([entry])).text.render(backend)


def get_value(data, key_path, fallbacks):
    keys = key_path.split(".")
    current = data
    for key in keys:
        if isinstance(current, dict) and key in current:
            current = current[key]
        else:
            return fallbacks.get(key_path, ""), False

    if current is None or (isinstance(current, str) and not current.strip()):
        return fallbacks.get(key_path, ""), False

    return current, isinstance(current, list)



def format_learning_outcomes(course_title, para, fallback_values):
    with open("classifications.yml", "r") as f:
        data = yaml.safe_load(f)
        grouped_results = {}
        outcomes = data.get("classifications", {}).get("learning_outcomes", {})
        for category, items in outcomes.items():
            for item in items:
                if isinstance(item.get("courses"), list) and course_title in item["courses"]:
                    grouped_results.setdefault(category, []).append(item["description"])

    if not grouped_results:
        raise RuntimeError("Cannot find learning outcomes.")

    category_labels = {
        "knowledge_and_understanding": "Knowledge and Understanding",
        "applying_knowledge_and_understanding": "Applying knowledge and understanding",
        "ability_to_make_judgments": "Ability to make judgments",
        "communication_skills": "Communication skills",
        "learning_skills": "Learning skills",
    }

    para.clear()
    first_written = False
    for category, descriptions in grouped_results.items():
        category_name = category_labels.get(category, category.replace("_", " ").title())
        category_para = para.insert_paragraph_before(category_name)
        category_para.style = "FirstLearningOutcome" if not first_written else "LearningOutcome"
        first_written = True
        for desc in descriptions:
            para.insert_paragraph_before(str(desc)).style = "Item"


def format_specific_educational_objectives(course_title, value):
    with open("classifications.yml", "r") as f:
        data = yaml.safe_load(f)
        course_type = next(
            (item["description"] for item in data["classifications"]["course_types"]
             if course_title in item.get("courses", [])),
            None
        )

        if course_type == None:
            raise RuntimeError("Unknown course type.")
        
        subject_area = next(
            (item["description"] for item in data["classifications"]["subject_area"]
             if course_title in item.get("courses", [])),
            None
        )

        if subject_area == None:
            raise RuntimeError("Unknown subject area type.")
        
    first_sentence = f'This course belongs to the type "{course_type}" and the subject area is "{subject_area}".'
    return [first_sentence, "", str(value)]


def replace_placeholders(paragraphs, data, bib_path):
    fallback_values = {
        "course.prerequisites": "There are no prerequisites for this course.",
        "course.course_page": "The course page will be made available on the Microsoft Teams class for this course or on https://ole.unibz.it, as communicated by the lecturer. Additional materials can also be found in the university's Reserve Collection at https://www.unibz.it/en/services/library/new-rc/.",
        "course.office_hours": "Office hours by appointment.",
        "module1.office_hours": "Office hours by appointment.",
        "module2.office_hours": "Office hours by appointment.",
        "course.supplementary_readings": "If supplementary readings are required, they will be communicated during class by the lecturer.",
        "course.software": "If the use of specific software is required, it will be communicated during class by the lecturer.",
    }

    name_keys_with_emails = {
        "course.lecturer.name": "course.lecturer.email",
        "course.teaching_assistant.name": "course.teaching_assistant.email",
        "module1.lecturer.name": "module1.lecturer.email",
        "module1.teaching_assistant.name": "module1.teaching_assistant.email",
        "module2.lecturer.name": "module2.lecturer.email",
        "module2.teaching_assistant.name": "module2.teaching_assistant.email",
    }

    office_hour_keys_with_offices = {
        "course.lecturer.office_hours": "course.lecturer.office",
        "module1.lecturer.office_hours": "module1.lecturer.office",
        "module2.lecturer.office_hours": "module2.lecturer.office",
        "course.teaching_assistant.office_hours": "course.teaching_assistant.office",
        "module1.teaching_assistant.office_hours": "module1.teaching_assistant.office",
        "module2.teaching_assistant.office_hours": "module2.teaching_assistant.office",
    }

    pattern = re.compile(r"\{\{([^}]+)\}\}")
    for para in paragraphs:
        matches = pattern.findall(para.text)
        if not matches:
            continue

        new_text = para.text
        replaced_list = False

        for match in matches:
            key = match.strip()
            value, is_list = get_value(data, key, fallback_values)

            if key in name_keys_with_emails:
                email_key = name_keys_with_emails[key]
                email_value, _ = get_value(data, email_key, fallback_values)
                if value and email_value:
                    if isinstance(value, list) and isinstance(email_value, list):
                        value = [
                            f"{n} ({e})"
                            for n, e in zip(value, email_value)
                            if n and e and str(n).strip() and str(e).strip()
                        ]
                    elif isinstance(value, str) and isinstance(email_value, str):
                        if value.strip() and email_value.strip():
                            value = f"{value} ({email_value})"

            if key in office_hour_keys_with_offices:
                office_key = office_hour_keys_with_offices[key]
                office_value, _ = get_value(data, office_key, fallback_values)
                if (
                    isinstance(value, str)
                    and value.strip()
                    and isinstance(office_value, str)
                    and office_value.strip()
                ):
                    value = f"Office {office_value}, {value[:1] + value[1:]}"

            if key == "course.modular":
                value = "Yes" if value else "No"

            if key == "course.specific_educational_objectives":
                course_title, _ = get_value(data, "course.title", fallback_values)
                paragraph_lines = format_specific_educational_objectives(course_title, value)
                para.clear()
                for line in paragraph_lines:
                    para.insert_paragraph_before(line)
                replaced_list = True
                break

            if key == "course.learning_outcomes":
                course_title, _ = get_value(data, "course.title", fallback_values)
                format_learning_outcomes(course_title, para, fallback_values)
                replaced_list = True
                break

            if isinstance(value, list):
                para.clear()
                for item in value:
                    if isinstance(item, str) and item.startswith("@"):
                        item = get_bib_entry(bib_path, item[1:])
                    para.insert_paragraph_before(str(item)).style = "Item"
                replaced_list = True
                break

            if isinstance(value, str) and value.startswith("@"):
                value = get_bib_entry(bib_path, value[1:])

            new_text = new_text.replace(f"{{{{{match}}}}}", str(value or "/"))

        if not replaced_list:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = new_text


def replace_in_tables(tables, data, bib_path):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders(cell.paragraphs, data, bib_path)
                if cell.paragraphs and not cell.paragraphs[-1].text.strip():
                    cell._element.remove(cell.paragraphs[-1]._element)


def convert_with_libreoffice(input_docx, output_pdf):
    try:
        subprocess.run(
            [
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "--headless",
                "--convert-to",
                "pdf",
                input_docx,
            ],
            check=True,
        )
        temp_pdf = os.path.splitext(os.path.basename(input_docx))[0] + ".pdf"
        os.rename(temp_pdf, output_pdf)
    except subprocess.CalledProcessError as e:
        print(f"Conversion failed: {e}")


def main():
    parser = argparse.ArgumentParser(description="Fill .docx template with YAML data.")
    parser.add_argument("template", help="Path to the .docx template file")
    parser.add_argument("output", help="Path to save the filled .docx file")
    parser.add_argument("yaml_files", nargs="+", help="YAML file(s) containing data")
    args = parser.parse_args()

    data = load_yaml(args.yaml_files)
    doc = Document(args.template)
    bib_path = "./references.bib"

    replace_placeholders(doc.paragraphs, data, bib_path)
    replace_in_tables(doc.tables, data, bib_path)
    doc.save(args.output)
    convert_with_libreoffice(args.output, args.output.replace(".docx", ".pdf"))


if __name__ == "__main__":
    main()